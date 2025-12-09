"""
Document parsing service for extracting pricing information from files
"""
import json
import re
import os
from typing import Dict, List, Optional
import pandas as pd
import pdfplumber
from docx import Document
from openpyxl import load_workbook
import logging
from app.config import settings
from app.services.ai_service import ai_service

logger = logging.getLogger(__name__)


class DocumentParserService:
    """Service for parsing documents and extracting pricing information"""
    
    def __init__(self):
        self.ai_service = ai_service
    
    def parse_document(self, file_path: str, file_extension: str) -> Dict:
        """
        Parse a document and extract pricing information
        
        Returns:
            {
                "success": bool,
                "products": List[Dict],
                "error": str or None
            }
        """
        try:
            # Extract text based on file type
            text_content = self._extract_text(file_path, file_extension)
            
            if not text_content:
                return {
                    "success": False,
                    "products": [],
                    "error": "Could not extract text from document"
                }
            
            # Use AI to extract pricing information
            products = self._extract_pricing_with_ai(text_content)
            
            return {
                "success": True,
                "products": products,
                "error": None
            }
            
        except Exception as e:
            logger.error(f"Error parsing document {file_path}: {e}")
            return {
                "success": False,
                "products": [],
                "error": str(e)
            }
    
    def _extract_text(self, file_path: str, file_extension: str) -> str:
        """Extract text from document based on file type"""
        ext = file_extension.lower()
        
        if ext == ".csv":
            return self._extract_from_csv(file_path)
        elif ext in [".xlsx", ".xls"]:
            return self._extract_from_excel(file_path)
        elif ext == ".pdf":
            return self._extract_from_pdf(file_path)
        elif ext in [".docx", ".doc"]:
            return self._extract_from_word(file_path)
        else:
            raise ValueError(f"Unsupported file type: {ext}")
    
    def _extract_from_csv(self, file_path: str) -> str:
        """Extract text from CSV file"""
        try:
            df = pd.read_csv(file_path)
            # Convert DataFrame to text representation
            text = df.to_string(index=False)
            # Also get column names and sample rows as context
            text += "\n\nColumns: " + ", ".join(df.columns.tolist())
            return text
        except Exception as e:
            logger.error(f"Error reading CSV: {e}")
            return ""
    
    def _extract_from_excel(self, file_path: str) -> str:
        """Extract text from Excel file"""
        try:
            # Try reading with pandas first
            df = pd.read_excel(file_path, engine='openpyxl')
            text = df.to_string(index=False)
            text += "\n\nColumns: " + ", ".join(df.columns.tolist())
            return text
        except Exception as e:
            logger.error(f"Error reading Excel: {e}")
            # Fallback: try reading raw with openpyxl
            try:
                wb = load_workbook(file_path)
                text = ""
                for sheet in wb.worksheets:
                    text += f"\nSheet: {sheet.title}\n"
                    for row in sheet.iter_rows(values_only=True):
                        text += " | ".join(str(cell) if cell else "" for cell in row) + "\n"
                return text
            except Exception as e2:
                logger.error(f"Error reading Excel with openpyxl: {e2}")
                return ""
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        try:
            text = ""
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            return text
        except Exception as e:
            logger.error(f"Error reading PDF: {e}")
            return ""
    
    def _extract_from_word(self, file_path: str) -> str:
        """Extract text from Word document"""
        try:
            doc = Document(file_path)
            text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
            
            # Also extract tables
            for table in doc.tables:
                text += "\n\nTable:\n"
                for row in table.rows:
                    text += " | ".join(cell.text for cell in row.cells) + "\n"
            
            return text
        except Exception as e:
            logger.error(f"Error reading Word document: {e}")
            return ""
    
    def _extract_pricing_with_ai(self, text_content: str) -> List[Dict]:
        """Use AI to extract pricing information from text"""
        if not settings.OPENAI_API_KEY:
            logger.warning("OpenAI API key not configured, cannot extract pricing with AI")
            return []
        
        try:
            prompt = f"""Extract product pricing information from this document.

Document Content:
{text_content}

Please extract all products with their pricing information. For each product, extract:
1. Product name
2. Pricing formula (e.g., "base_price + (area * rate)", "area * 0.05", "if area > 1000 then rate * 1.1 else rate")
3. Base price (if applicable)
4. Rate per square inch (if applicable)
5. Minimum size in square inches
6. Maximum size in square inches
7. Product description (if available)

The pricing formula should be a Python-like expression that can calculate price from:
- area: area in square inches (length * width)
- base_price: base price for the product
- rate: price per square inch

Return a JSON array of products in this format:
[
    {{
        "name": "Product Name",
        "description": "Product description",
        "base_price": 10.0 or null,
        "rate": 0.05 or null,
        "pricing_formula": "base_price + (area * rate)" or "area * 0.05",
        "min_size_sq_in": 100,
        "max_size_sq_in": 10000
    }}
]

If pricing information is not available in a formula format, try to derive it from the data provided.
Return an empty array if no products are found."""

            # Use OpenAI to extract information
            from openai import OpenAI
            client = OpenAI(api_key=settings.OPENAI_API_KEY)
            
            response = client.chat.completions.create(
                model=settings.AI_MODEL,
                messages=[
                    {"role": "system", "content": "You are a helpful assistant that extracts structured pricing information from documents. Always return valid JSON."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.1,
                max_tokens=2000
            )
            
            content = response.choices[0].message.content.strip()
            
            # Remove markdown code blocks if present
            content = re.sub(r'```json\n?', '', content)
            content = re.sub(r'```\n?', '', content)
            
            products = json.loads(content)
            
            # Validate and normalize products
            normalized_products = []
            for product in products:
                normalized = self._normalize_product(product)
                if normalized:
                    normalized_products.append(normalized)
            
            return normalized_products
            
        except Exception as e:
            logger.error(f"Error extracting pricing with AI: {e}")
            return []
    
    def _normalize_product(self, product: Dict) -> Optional[Dict]:
        """Normalize product data and validate"""
        try:
            name = product.get("name", "").strip()
            if not name:
                return None
            
            # Get pricing formula
            formula = product.get("pricing_formula", "")
            if not formula:
                # Try to construct formula from base_price and rate
                base_price = product.get("base_price", 0) or 0
                rate = product.get("rate", 0) or 0
                
                if base_price > 0 and rate > 0:
                    formula = f"{base_price} + (area * {rate})"
                elif rate > 0:
                    formula = f"area * {rate}"
                elif base_price > 0:
                    formula = str(base_price)
                else:
                    # Default formula
                    formula = "area * 0.05"
            
            return {
                "name": name,
                "description": product.get("description", ""),
                "pricing_formula": formula,
                "base_price": product.get("base_price", 0) or 0,
                "rate": product.get("rate", 0) or 0,
                "min_size_sq_in": float(product.get("min_size_sq_in", 100)),
                "max_size_sq_in": float(product.get("max_size_sq_in", 10000))
            }
            
        except Exception as e:
            logger.error(f"Error normalizing product: {e}")
            return None


# Singleton instance
document_parser_service = DocumentParserService()



